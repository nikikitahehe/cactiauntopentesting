from docx import Document
from docx.shared import Pt, Cm, Inches
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from datetime import date
from docx.oxml import OxmlElement
from docx.oxml.ns import nsdecls
from docx.oxml.ns import qn
import subprocess

def add_header(document):
    # Add Header
    header_section = document.sections[0]
    header = header_section.header
    header_paragraph = header.paragraphs[0]
    header_paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
    
    # Add the logo
    logo_path = "Logo Horizontal.png"
    logo_height = Cm(1.5)
    run = header_paragraph.add_run()
    run.add_picture(logo_path, height=logo_height)
    
    # Add space after header
    header.add_paragraph()

def create_cover_page(document):
    add_header(document)
    
    # Add the logo image
    document.add_picture('tri.png', width=Inches(2))
    document.paragraphs[-1].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER  # Center align the image
    
    # Add spacer
    document.add_paragraph()
    document.paragraphs[-1].add_run().add_break()
    
    while len(document.paragraphs) % 4 != 0:
        document.add_paragraph()
    
    title = """Sample Penetration Test Report
    Example Company"""
    document.add_heading(title, level=0).alignment = WD_PARAGRAPH_ALIGNMENT.CENTER  # Center align the text
    
    # Add spacer
    document.add_paragraph()
    document.paragraphs[-1].add_run().add_break()
    
    while len(document.paragraphs) % 14 != 0:
        document.add_paragraph()
    
    # Add details
    text = f"""
    Company: Customer Name
    Date: {date.today().strftime('%d %B %Y')}
    Version 1.0"""
    document.add_paragraph(text)

def add_pendahuluan(document):
    document.add_heading("Pendahuluan", level=2)
    pendahuluan_paragraph = document.add_paragraph()
    pendahuluan_paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY  # Rata kanan-kiri
    pendahuluan_text = (
        "Laporan ini disusun sebagai hasil pengujian penetrasi terhadap CVE-2022-46169, "
        "sebuah kerentanan yang ditemukan dalam perangkat lunak Cacti. Kerentanan ini "
        "memungkinkan serangan tanpa autentikasi untuk melakukan eksekusi perintah sistem "
        "secara sewenang-wenang pada server yang menjalankan Cacti. Dalam laporan ini, kami "
        "akan memberikan detail mengenai temuan kerentanan CVE-2022-46169 yang kami identifikasi "
        "selama penilaian. Kami akan menjelaskan dengan rinci potensi dampak dan risiko yang "
        "terkait dengan kerentanan ini, serta memberikan rekomendasi tindakan untuk memperbaiki "
        "kerentanan tersebut dan meningkatkan keamanan sistem secara menyeluruh."
    )
    pendahuluan_paragraph.add_run(pendahuluan_text)

def add_ruang_lingkup(document):
    document.add_heading("Ruang Lingkup", level=2)
    ruang_lingkup_paragraph = document.add_paragraph()
    ruang_lingkup_paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY  # Rata kanan-kiri
    ruang_lingkup_text = (
        "Evaluasi keamanan sistem informasi pada Cacti dilakukan di lingkungan produksi "
        "dengan melakukan upaya peretasan berdasarkan kerentanan yang ditemukan dan "
        "alamat IP yang diuji adalah sebagai berikut:"
    )
    ruang_lingkup_paragraph.add_run(ruang_lingkup_text)
    
    # Add the hosts and IP addresses
    with open("data.txt", "r") as data_file:
        data_content = data_file.read()
    document.add_paragraph(data_content)

def add_metodologi(document):
    document.add_heading("Metodologi", level=2)
    metodologi_paragraph = document.add_paragraph()
    metodologi_paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY # Rata kanan-kiri
    metodologi_text = (
        "Metodologi yang digunakan dalam pengujian penetrasi ini terdiri dari beberapa tahap "
        "yang sistematis untuk memastikan pengujian yang menyeluruh dan efektif. Tahap-tahap "
        "ini meliputi information gathering, vulnerability scanning, vulnerability analysis, "
        "vulnerability exploitation, recommendation and reporting. Metodologi ini dirancang untuk "
        "mengidentifikasi dan mengatasi potensi celah keamanan dalam sistem secara menyeluruh."
    )
    metodologi_paragraph.add_run(metodologi_text)
    vuln_devices = []
    vuln_devices_public = []
    with open("data.txt", 'r') as f_ip:
        ip_addresses = f_ip.readlines()

# Loop melalui setiap alamat IP dalam daftar
    for ip_address in ip_addresses:
        ip_address = ip_address.strip()
    # Buka file vuln_scan.txt untuk memeriksa laporan kerentanan
    with open("vuln_scan_public.txt", 'r') as f_vuln:
        lines = f_vuln.readlines()
    i = 0
    while i < len(lines):
        line = lines[i].strip()
        if "========== Vulnerability Scan Result for" in line:
            scanned_ip = line.split()[-1]
            if ip_address == scanned_ip:
                if i + 5 < len(lines):
                    status_line = lines[i+5].strip()
                    if "The target appears to be vulnerable. The target is Cacti version 1.2.22" in status_line:
                        vuln_devices.append((ip_address, "CVE-2022-46169 Vulnerability Detected"))
                    else:
                        vuln_devices.append((ip_address, "No Vulnerability Detected"))
                break
        i += 1

    with open("vuln_scan_public.txt", 'r') as f_vuln:
        lines = f_vuln.readlines()
    i = 0
    while i < len(lines):
        line = lines[i].strip()
        if "========== Vulnerability Scan Result for" in line:
            ip_address = line.split()[-1]
            if i + 5 < len(lines):
                status_line = lines[i+5].strip()
                if "The target appears to be vulnerable. The target is Cacti version 1.2.22" in status_line:
                    vuln_devices_public.append((ip_address, "CVE-2022-46169 Vulnerability Detected"))
                else:
                    vuln_devices_public.append((ip_address, "No CVE-2022-46169 Vulnerability Detected"))
        i += 1

# Adding Target Devices if either list has items
    if vuln_devices or vuln_devices_public:
        document.add_paragraph("Vulnerable Devices:")
    for device in vuln_devices_public:
        document.add_paragraph(f"- {device[0]}: {device[1]}")
def add_vulnerability_identification(document, output_file):
    document.add_heading("Identifikasi Kerentanan", level=2)
    identification_paragraph = document.add_paragraph()
    identification_paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY  # Rata kanan-kiri
    identification_text = (
        "Pemindaian atau pencarian menggunakan Shodan dengan kata kunci \"Cacti\" "
        "memungkinkan untuk menemukan sistem-sistem yang menggunakan aplikasi Cacti."
    )
    identification_paragraph.add_run(identification_text)

    # Adding content from nmap_results.txt
    with open(output_file, "r") as scan_file:
        scan_content = scan_file.read()
    document.add_paragraph(scan_content)
    
    scan_results = (
    "Shodan adalah mesin pencari untuk perangkat yang terhubung ke internet yang memberikan informasi tentang perangkat keras, "
    "perangkat lunak, konfigurasi, dan jenis layanan yang berjalan pada perangkat tersebut. Dengan menggunakan kata kunci \"Cacti\", "
    "Shodan akan mengembalikan hasil yang mencakup alamat IP dari sistem yang terdeteksi menggunakan aplikasi Cacti, port tempat "
    "aplikasi Cacti berjalan (umumnya pada port 80 untuk HTTP), nama penyedia jaringan atau organisasi yang terkait dengan alamat "
    "IP tersebut, serta tanggal dan waktu kapan data terakhir kali diambil atau sistem dipindai oleh Shodan. Contoh informasi ini "
    "menunjukkan bahwa beberapa alamat IP yang terhubung ke jaringan Biznet Networks dan PT Biznet Gio Nusantara menggunakan aplikasi "
    "Cacti pada port 80. Informasi ini memberikan pemahaman tentang di mana aplikasi Cacti diimplementasikan dan konteks infrastruktur "
    "jaringan yang digunakan."
    )

    document.add_paragraph(scan_results)

    
def add_vulnerability_scanning(document):
    document.add_heading("Vulnerability Scanning", level=2)
    scanning_paragraph = document.add_paragraph()
    scanning_paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY  # Rata kanan-kiri
    scanning_text = (
        "Vulnerability scanning dilakukan menggunakan perangkat lunak Metasploit. Metasploit adalah open-source, "
        "platform pengujian penetrasi berbasis Ruby yang memungkinkan pengguna untuk menulis, menguji, dan mengeksekusi "
        "kode eksploit. Sistem pengujian penetrasi atau pengujian pena bekerja dengan mensimulasikan serangan cyber untuk "
        "memeriksa kerentanan yang rentan. Dibawah ini menampilkan hasil dari pemindaian kerentanan yang ditemukan oleh Metasploit."
    )
    scanning_paragraph.add_run(scanning_text)
    
    # Add vulnerability scan results from file
    with open("vuln_scan_public.txt", "r") as vuln_file:
        vuln_content = vuln_file.read()
    document.add_paragraph(vuln_content)

def add_recommendations(document):
    document.add_heading("Recommendation", level=2)
    recommendations_paragraph = document.add_paragraph()
    recommendations_paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY  # Rata kanan-kiri
    recommendations_text = "Untuk mengurangi risiko dari CVE-2022-46169, disarankan untuk mengambil langkah-langkah berikut:"
    recommendations_paragraph.add_run(recommendations_text)

    # Add the recommendations list
    recommendations_list = [
        "Memperbarui Cacti ke versi terbaru yang tersedia.",
        "Menerapkan aturan firewall yang membatasi akses ke layanan Cacti.",
        "Melakukan evaluasi keamanan secara berkala dan pengujian penetrasi untuk mengidentifikasi dan mengatasi kerentanan.",
        "Menerapkan kebijakan sandi yang kuat dan menghindari penggunaan kredensial default.",
        "Rutin memperbarui perangkat untuk mengatasi masalah keamanan."
    ]
    for recommendation in recommendations_list:
        document.add_paragraph(recommendation, style='List Bullet')
    
    

def generate_report(output_file):
    # Create a new Word document
    document = Document()

    # Generate the cover page
    create_cover_page(document)

    # Add Pendahuluan
    add_pendahuluan(document)

    # Add Ruang Lingkup
    add_ruang_lingkup(document)

    # Add Metodologi
    add_metodologi(document)

    # Add Vulnerability Identification
    add_vulnerability_identification(document,output_file)

    # Add Vulnerability Scanning
    add_vulnerability_scanning(document)

    # Add Recommendations
    add_recommendations(document)

    # Save the document
    document.save("Penetration_Test_Report_Public.docx")

    # Convert to PDF
    subprocess.run(["unoconv", "-f", "pdf", "Penetration_Test_Report_Public.docx"], stderr=subprocess.DEVNULL)

    print("Penetration test report generated successfully.")

if __name__ == "__main__":
    generate_report()
